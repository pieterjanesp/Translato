from docx import Document
from app.handlers.base import BaseDocumentHandler

class DOCXHandler(BaseDocumentHandler):
    """Handler for DOCX files."""

    def extract_text(self, file_path: str) -> list[str]:
        """Extract text from DOCX file."""
        doc = Document(file_path)
        texts = [
            paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()
        ]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        texts.append(cell.text)
        return texts

    def replace_text(self, file_path: str, translations: dict[str, str], output_path: str) -> None:
        """Replace text in DOCX file with translations."""
        doc = Document(file_path)
        #Replace text in paragraphs
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                for original,translated in translations.items():
                    if original in run.text:
                        run.text = run.text.replace(original, translated)
        #Replace text in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            for original,translated in translations.items():
                                if original in run.text:
                                    run.text = run.text.replace(original, translated)
        # Save the modified document
        doc.save(output_path)